"""
Step 04: Dosage Extraction & Inference

從已找到替代物的同一篇論文全文中提取劑量資訊。
若無明確數字，則根據論文內容用 LLM 進行推論。

流程：
1. 讀取 step03 結果（已包含替代物資訊）
2. 對每篇有替代物的論文：
   a. 檢查是否有下載的 PDF/XML 全文
   b. 若有全文，用 LLM 從全文中提取劑量資訊
   c. 若無明確劑量，用 LLM 根據全文/摘要進行推論
3. 輸出包含劑量資訊的結果
"""

import argparse
import json
import re
from pathlib import Path
from typing import Any, Optional

# Use Windows system CA certificates instead of certifi (fixes SSL issues on some networks)
try:
    import truststore
    truststore.inject_into_ssl()
except ImportError:
    pass  # truststore not installed, use default certifi

from openai import OpenAI
from tqdm import tqdm

# Gemini support
try:
    import google.genai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    genai = None
    GEMINI_AVAILABLE = False

from config_loader import get_config
# Import downloader to dynamically fetch PDFs if they are missing
import sys
import os
try:
    from step03 import download_full_text
except ImportError:
    pass

# Load configuration
CONFIG = get_config()
OPENAI_API_KEY = CONFIG.get("openai_api_key", "")
GEMINI_API_KEY = CONFIG.get("gemini_api_key", "")
LLM_PROVIDER = CONFIG.get("default_settings", {}).get("llm_provider", "openai")
MODEL = CONFIG.get("default_settings", {}).get("openai_model", "gpt-4o-mini")
GEMINI_MODEL = CONFIG.get("default_settings", {}).get("gemini_model", "gemini-2.0-flash")
MAX_RETRIES = CONFIG.get("default_settings", {}).get("max_retries", 3)

# ============================================================
# Dosage Type Classification
# ============================================================

# Material dosage types (INPUT: how much reagent to use)
MATERIAL_DOSAGE_TYPES = {
    "molar_ratio", "eq", "equivalent", "equivalents",
    "wt%", "wt percent", "weight_percent", "weight_ratio",
    "vol%", "volume_percent", "volume_ratio",
    "mol%", "molar_percent",
    "catalyst_loading", "loading",
    "mass_ratio", "feed_ratio",
}

# Result dosage types (OUTPUT: product specification)
RESULT_DOSAGE_TYPES = {
    "concentration", "content", "yield",
    "degree_of_functionalization", "conversion",
    "capacity", "activity",
}


def is_material_dosage(dosage: dict) -> bool:
    """Check if a dosage entry is a material dosage (input) vs result dosage (output)."""
    ratio_type = (dosage.get("ratio_type") or "").lower().replace(" ", "_")
    value = (dosage.get("value") or "").lower()
    
    # Check ratio_type against known material types
    for mat_type in MATERIAL_DOSAGE_TYPES:
        if mat_type in ratio_type:
            return True
    
    # Check value for material dosage patterns (e.g., "3 eq", "10 wt%", "5 mol/L")
    material_patterns = [
        r'\d+\.?\d*\s*(eq|equiv)',
        r'\d+\.?\d*\s*(wt|vol|mol)\s*%',
        r'\d+\.?\d*\s*mg\s*/\s*(mL|L|g)',
        r'\d+\.?\d*\s*g\s*/\s*(mol|mmol)',
    ]
    for pattern in material_patterns:
        if re.search(pattern, value, re.IGNORECASE):
            return True
    
    return False


def is_synthesis_scale(dosage: dict) -> bool:
    """Return True if this dosage entry represents a total lab batch size
    (e.g., '>55 g synthesized') or a production/industrial scale quantity
    (e.g., '291.6 kg catalyst per stage', '16700 t/year feedstock'),
    rather than a functional composition ratio.
    Such values are useless for functional-equivalence dosage comparison.
    """
    ratio_type = (dosage.get("ratio_type") or "").lower()
    value = (dosage.get("value") or "").lower()
    context = (dosage.get("context") or "").lower()

    # ── Guard 1: production/industrial throughput units (always discard) ──────
    # e.g. "16700.2 t/year", "4.2 billion liters", "20873 tonnes/year"
    if re.search(r'\d[\d.,]*\s*(million|billion|thousand)?\s*(t|tonne|ton|liters?|litres?)'  # volume/mass
                 r'\s*(/|per)\s*(year|yr|annum|day)',
                 value, re.IGNORECASE):
        return True
    if re.search(r'\d[\d.,]*\s*(million|billion)\s*(liters?|litres?|kg|g|t\b)', value, re.IGNORECASE):
        return True

    # ── Guard 2: catalyst_loading or mass_ratio with absolute-mass units ──────
    if ratio_type not in ("mass_ratio", "catalyst_loading"):
        return False

    # If value contains a ratio separator or relative unit, it IS compositional — keep it
    if re.search(r'[:/]|wt%|vol%|mol%|\beq\b|\bequiv\b|\bwt\b|\bvol\b|\bmol\b', value):
        return False

    # Absolute mass (g, mg, kg) with synthesis-scale context keywords → discard
    is_absolute_mass = bool(re.search(r'\d+\.?\d*\s*(>|<|~)?\s*(kg|g|mg)\b', value))
    synthesis_keywords = (
        "synthesis scale", "synthesized", "batch", "prepared", "scale up",
        "total yield", "total amount", "lab scale", "gram scale",
        "milligram scale", "reaction scale", "synthesis of",
        # industrial process contexts
        "per stage", "per tray", "reactor", "distillation", "column",
        "plant", "production unit",
    )
    if is_absolute_mass and any(kw in context for kw in synthesis_keywords):
        return True

    return False


def is_impurity_concentration(dosage: dict) -> bool:
    """Return True if this dosage entry measures an impurity / contaminant /
    leached species within a material, or a macroscopic production/consumption
    statistic, NOT the functional dosage of the alternative material itself.
    These pollute explicit_dosages and should be discarded.
    """
    ratio_type = (dosage.get("ratio_type") or "").lower()
    context = (dosage.get("context") or "").lower()
    material = (dosage.get("material") or "").lower()
    value = (dosage.get("value") or "").lower()

    if ratio_type != "concentration":
        return False

    # Context signals impurity / contamination measurement
    impurity_context_keywords = (
        "found in", "concentration in", "detected in", "leaching",
        "migration", "leakage", "impurity", "contamination",
        "additive", "extracted from", "measured in", "content in",
        "mean concentration",
    )
    if any(kw in context for kw in impurity_context_keywords):
        return True

    # Context signals global/national production or consumption statistics
    production_stat_keywords = (
        "annual consumption", "annual production", "annual output",
        "global consumption", "global production", "worldwide consumption",
        "per year", "per annum", "raising the octane", "octane number",
        "for raising",
    )
    if any(kw in context for kw in production_stat_keywords):
        return True

    # Value itself is a production-scale volume (already caught by is_synthesis_scale
    # for other ratio_types; duplicate check here for concentration type)
    if re.search(r'\d[\d.,]*\s*(million|billion|thousand)?\s*(t|tonne|ton|liters?|litres?)'
                 r'\s*(/|per)\s*(year|yr|annum)',
                 value, re.IGNORECASE):
        return True

    # Material name is a metal element / ion (not the alternative polymer)
    metal_pattern = (
        r'^(barium|ba|zinc|zn|lead|pb|cadmium|cd|tin|sn|chromium|cr|'
        r'mercury|hg|arsenic|as|iron|fe|copper|cu|nickel|ni|cobalt|co)\b'
    )
    if re.match(metal_pattern, material, re.IGNORECASE):
        return True

    return False


def filter_functional_dosages(explicit_dosages: list[dict] | None) -> list[dict]:
    """Remove synthesis-scale and impurity-concentration entries from a
    dosages list, returning only entries that represent true functional
    composition / usage ratios of the alternative material.
    """
    if not explicit_dosages:
        return []
    return [
        d for d in explicit_dosages
        if not is_synthesis_scale(d) and not is_impurity_concentration(d)
    ]


def has_material_dosage(explicit_dosages: list[dict] | None) -> bool:
    """Check if any of the explicit dosages is a material dosage."""
    if not explicit_dosages:
        return False
    return any(is_material_dosage(d) for d in explicit_dosages)


def classify_extraction_completeness(extraction_result: dict) -> str:
    """
    Classify extraction result as:
    - 'complete': has material dosage (input quantities)
    - 'partial_result_only': has only result/output dosages
    - 'insufficient': no dosages found

    Filters out synthesis-scale values and impurity concentrations before
    classifying, so those false positives no longer count as 'complete'.
    """
    if not extraction_result.get("dosage_found"):
        return "insufficient"

    raw_dosages = extraction_result.get("explicit_dosages", [])
    if not raw_dosages:
        return "insufficient"

    dosages = filter_functional_dosages(raw_dosages)
    if not dosages:
        return "insufficient"

    if has_material_dosage(dosages):
        return "complete"
    else:
        return "partial_result_only"

# Validate API keys based on provider
if LLM_PROVIDER == "gemini":
    if not GEMINI_API_KEY:
        print("[ERROR] Gemini API key not found! Set GEMINI_API_KEY in .env")
        exit(1)
    if not GEMINI_AVAILABLE:
        print("[ERROR] google-genai not installed. Run: pip install google-genai")
        exit(1)
else:
    if not OPENAI_API_KEY:
        print("[ERROR] OpenAI API key not found in config!")
        exit(1)


# ============================================================
# LLM Client Wrapper (supports OpenAI and Gemini)
# ============================================================

class LLMClient:
    """Unified LLM client supporting OpenAI and Gemini."""
    
    def __init__(self, provider: str = None):
        self.provider = (provider or LLM_PROVIDER).lower()
        self.usage = {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0}
        
        if self.provider == "gemini":
            self.gemini_client = genai.Client(api_key=GEMINI_API_KEY)
            self._openai_client = None
            print(f"[INFO] Using Gemini provider with model: {GEMINI_MODEL}")
        else:
            self._openai_client = OpenAI(api_key=OPENAI_API_KEY, timeout=60.0)
            print(f"[INFO] Using OpenAI provider with model: {MODEL}")
    def chat_completion(
        self,
        messages: list[dict],
        temperature: float = 0.0,
        max_tokens: int = 1000,
        json_mode: bool = True,
        model: str = None
    ) -> str:
        """
        Send a chat completion request.
        Returns the response text content.
        """
        if self.provider == "gemini":
            return self._gemini_completion(messages, temperature, max_tokens, json_mode, model)
        else:
            return self._openai_completion(messages, temperature, max_tokens, json_mode, model)
    
    def _openai_completion(
        self,
        messages: list[dict],
        temperature: float,
        max_tokens: int,
        json_mode: bool,
        model: str
    ) -> str:
        model = model or MODEL
        kwargs = {
            "model": model,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": max_tokens,
        }
        if json_mode:
            kwargs["response_format"] = {"type": "json_object"}
        
        response = self._openai_client.chat.completions.create(**kwargs)
        if hasattr(response, "usage") and response.usage:
            self.usage["prompt_tokens"] += getattr(response.usage, "prompt_tokens", 0) or 0
            self.usage["completion_tokens"] += getattr(response.usage, "completion_tokens", 0) or 0
            self.usage["total_tokens"] += getattr(response.usage, "total_tokens", 0) or 0
        return response.choices[0].message.content or ""
    
    def _gemini_completion(
        self,
        messages: list[dict],
        temperature: float,
        max_tokens: int,
        json_mode: bool,
        model: str
    ) -> str:
        # Convert OpenAI-style messages to Gemini format
        prompt_parts = []
        for msg in messages:
            role = msg.get("role", "user")
            content = msg.get("content", "")
            if role == "system":
                prompt_parts.insert(0, f"[System]: {content}\n\n")
            else:
                prompt_parts.append(content)
        
        full_prompt = "".join(prompt_parts)
        
        # Determine which model to use
        gemini_model = model or GEMINI_MODEL
        if model == "gpt-4o-mini" or "fast" in (model or "").lower():
            gemini_model = "gemini-2.0-flash"
        
        response = self.gemini_client.models.generate_content(
            model=gemini_model,
            contents=full_prompt,
            config=genai.types.GenerateContentConfig(
                temperature=temperature,
                max_output_tokens=max_tokens,
                response_mime_type="application/json" if json_mode else None
            )
        )
        
        if hasattr(response, "usage_metadata") and response.usage_metadata:
            self.usage["prompt_tokens"] += getattr(response.usage_metadata, "prompt_token_count", 0) or 0
            self.usage["completion_tokens"] += getattr(response.usage_metadata, "candidates_token_count", 0) or 0
            self.usage["total_tokens"] += getattr(response.usage_metadata, "total_token_count", 0) or 0
            
        raw = response.text or ""
        # Handle potential markdown code blocks
        if raw.startswith("```"):
            lines = raw.split("\n")
            raw = "\n".join(lines[1:-1]) if lines[-1].strip() == "```" else "\n".join(lines[1:])
        return raw


# ============================================================
# Fulltext Reading Utilities
# ============================================================

def is_valid_pdf(file_path: Path) -> bool:
    """Check if a file is a valid PDF by verifying magic bytes."""
    try:
        with open(file_path, "rb") as f:
            header = f.read(8)
            return header.startswith(b"%PDF")
    except Exception:
        return False


def read_pdf_text(pdf_path: Path) -> str | None:
    """Extract text from PDF using PyMuPDF (pymupdf) or pdfplumber."""
    text = ""
    
    # Validate PDF magic bytes first
    if not is_valid_pdf(pdf_path):
        safe_name = str(pdf_path.name).encode(sys.stdout.encoding or 'ascii', errors='replace').decode(sys.stdout.encoding or 'ascii')
        print(f"    [WARN] Invalid PDF (not a real PDF file): {safe_name}. Deleting it.")
        try:
            pdf_path.unlink()
        except Exception as e:
            print(f"    [WARN] Failed to delete {safe_name}: {e}")
        return None
    
    # Try PyMuPDF first (faster)
    try:
        import pymupdf  # PyMuPDF
        doc = pymupdf.open(pdf_path)
        for page in doc:
            text += page.get_text()
        doc.close()
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception as e:
        safe_name = str(pdf_path.name).encode(sys.stdout.encoding or 'ascii', errors='replace').decode(sys.stdout.encoding or 'ascii')
        print(f"    [WARN] PyMuPDF failed for {safe_name}: {e}".encode(sys.stdout.encoding or 'ascii', errors='replace').decode(sys.stdout.encoding or 'ascii'))

    # Fallback to pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text
    except ImportError:
        pass
    except Exception as e:
        safe_name = str(pdf_path.name).encode(sys.stdout.encoding or 'ascii', errors='replace').decode(sys.stdout.encoding or 'ascii')
        print(f"    [WARN] pdfplumber failed for {safe_name}: {e}".encode(sys.stdout.encoding or 'ascii', errors='replace').decode(sys.stdout.encoding or 'ascii'))

    return text


def read_xml_text(xml_path: Path) -> str:
    """Extract text content from XML (Elsevier/JATS format)."""
    try:
        import xml.etree.ElementTree as ET
        tree = ET.parse(xml_path)
        root = tree.getroot()
        
        # Extract all text content
        texts: list[str] = []
        for elem in root.iter():
            if elem.text:
                texts.append(elem.text.strip())
            if elem.tail:
                texts.append(elem.tail.strip())
        
        return " ".join(t for t in texts if t)
    except Exception as e:
        print(f"    [WARN] XML parsing failed for {xml_path.name}: {e}")
        return ""


def read_docx_text(docx_path: Path) -> str:
    """Extract plain text from a .docx (Word) supplementary file."""
    try:
        from docx import Document
        doc = Document(docx_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"    [WARN] DOCX parsing failed for {docx_path.name}: {e}")
        return ""


def find_fulltext_file(doi: str, research_pdf_dir: Path) -> Optional[Path]:
    """Find the fulltext file (PDF or XML) for a given DOI."""
    if not doi or not research_pdf_dir.exists():
        return None
    
    # Normalize DOI for filename matching
    doi_normalized = doi.replace("/", "_").replace(".", "_").lower()
    doi_simple = doi.replace("/", "_").lower()
    
    for file in research_pdf_dir.iterdir():
        if file.is_file():
            filename_lower = file.name.lower()
            # Check if DOI is in the filename (exact match)
            if doi_simple in filename_lower or doi_normalized in filename_lower.replace(".", "_"):
                return file
    
    # Try more flexible matching - but require UNIQUE identifier part to match
    # DOI format: 10.XXXX/suffix - we need the suffix to match
    doi_suffix = doi.split("/")[-1].lower() if "/" in doi else doi.lower()
    
    for file in research_pdf_dir.iterdir():
        if file.is_file():
            filename_lower = file.name.lower()
            # The unique suffix part must be in the filename
            if doi_suffix in filename_lower:
                return file
    
    return None


def find_esi_files(doi: str, research_pdf_dir: Path) -> list[Path]:
    """
    Find ESI (Electronic Supplementary Information) files for a given DOI.
    
    ESI files typically have naming patterns like:
    - CSSC-18-e202402051-s001.pdf  (journal abbreviation + article ID + s00X)
    - xxx_ESI_yyy.pdf (with _ESI_ marker)
    - DOI-based: 10.1002_cssc.202402051_ESI.pdf
    - Wiley: cssc202402051-sup-0001-misc_information.pdf
    
    Returns list of ESI file paths found.
    """
    if not doi or not research_pdf_dir.exists():
        return []
    
    esi_files: list[Path] = []
    
    # Extract identifiers from DOI for matching
    # e.g., "10.1002/cssc.202402051" -> ["cssc", "202402051", "cssc.202402051"]
    doi_lower = doi.lower()
    doi_suffix = doi.split("/")[-1].lower() if "/" in doi else doi_lower  # "cssc.202402051"
    
    # Extract article ID pattern (e.g., "202402051" or "e202402051")
    article_id_match = re.search(r'[a-z]*(\d{6,})', doi_suffix)
    article_id = article_id_match.group(0) if article_id_match else ""
    
    # Extract journal abbreviation from DOI suffix.
    # Strategy: take the longest purely-alphabetic segment (≥2 chars) among all dot-separated parts.
    # This handles:
    #   - "cssc.202402051"          → "cssc"
    #   - "j.radphyschem.2017.07.008" → "radphyschem"  (Elsevier: "j." is publisher prefix)
    #   - "acs.jchemed.1c00001"     → "jchemed"
    doi_parts = doi_suffix.split(".")
    alpha_parts = [p for p in doi_parts if p.isalpha() and len(p) >= 2]
    journal_abbrev = max(alpha_parts, key=len) if alpha_parts else ""
    
    for file in research_pdf_dir.iterdir():
        if not file.is_file():
            continue
        if file.suffix.lower() not in (".pdf", ".xml", ".docx"):
            continue
        fname = file.name.lower()
        
        # Skip main article files (already handled by find_fulltext_file)
        # Main files typically have the DOI directly in name WITHOUT ESI markers
        doi_simple = doi.replace("/", "_").lower()
        # Only skip if it has DOI but NO ESI markers
        has_esi_marker = (
            "_esi" in fname or 
            "supporting" in fname or 
            "supplementary" in fname or 
            "-sup-" in fname or  # Wiley pattern: -sup-0001
            re.search(r'-s\d{3}', fname)
        )
        if doi_simple in fname and not has_esi_marker:
            continue
        
        # Check for ESI patterns
        is_esi = False
        
        # Pattern 1: Contains "-s001", "-s002", etc. (supplementary file suffix)
        if re.search(r'-s\d{3}', fname):
            # Must also match journal/article identifier
            if journal_abbrev and journal_abbrev in fname:
                is_esi = True
            elif article_id and article_id in fname:
                is_esi = True
            elif doi_simple in fname:
                is_esi = True
        
        # Pattern 2: Wiley pattern "-sup-0001", "-sup-0002", etc.
        if re.search(r'-sup-\d{4}', fname):
            # Must match article_id or journal abbreviation
            if article_id and article_id in fname:
                is_esi = True
            elif journal_abbrev and journal_abbrev in fname:
                is_esi = True
        
        # Pattern 3: Contains "_ESI" or "_esi" (our naming convention)
        if "_esi" in fname:
            # Must match DOI identifier or article_id
            if doi_simple in fname:
                is_esi = True
            elif article_id and article_id in fname:
                is_esi = True
            elif journal_abbrev and journal_abbrev in fname:
                is_esi = True
        
        # Pattern 4: Contains "supporting" or "supplementary" in name
        if "supporting" in fname or "supplementary" in fname:
            if article_id and article_id in fname:
                is_esi = True
            elif doi_simple in fname:
                is_esi = True
        
        if is_esi:
            esi_files.append(file)
    
    return esi_files


def get_fulltext(doi: str, research_pdf_dir: Path) -> tuple[str, str]:
    """
    Get fulltext content for a paper, including ESI if available.
    Returns (text, source) where source is 'pdf', 'xml', 'pdf+esi', 'xml+esi', or 'none'.
    """
    fulltext_file = find_fulltext_file(doi, research_pdf_dir)
    main_text = ""
    source = "none"
    
    # Read main article
    if fulltext_file:
        if fulltext_file.suffix.lower() == ".pdf":
            text = read_pdf_text(fulltext_file)
            if text is None:
                return "", "deleted_pdf"
            main_text = text
            source = "pdf" if text else "none"
        elif fulltext_file.suffix.lower() == ".xml":
            main_text = read_xml_text(fulltext_file)
            source = "xml" if main_text else "none"
    
    # Find and read ESI files
    esi_files = find_esi_files(doi, research_pdf_dir)
    esi_texts: list[str] = []
    
    for esi_file in esi_files:
        esi_text = ""
        if esi_file.suffix.lower() == ".pdf":
            esi_text = read_pdf_text(esi_file) or ""
        elif esi_file.suffix.lower() == ".xml":
            esi_text = read_xml_text(esi_file) or ""
        elif esi_file.suffix.lower() == ".docx":
            esi_text = read_docx_text(esi_file) or ""
        
        if esi_text:
            esi_texts.append(f"\n\n=== SUPPLEMENTARY INFORMATION ({esi_file.name}) ===\n\n{esi_text}")
            print(f"    [INFO] Found ESI: {esi_file.name} ({len(esi_text)} chars)")
    
    # Combine main text with ESI
    if esi_texts:
        combined_text = main_text + "".join(esi_texts)
        source = f"{source}+esi" if source != "none" else "esi"
        return combined_text, source
    
    return main_text, source


# ============================================================
# LLM Prompts
# ============================================================

DOSAGE_EXTRACTION_PROMPT = """# Role & Objective
You are an expert Data Extraction AI specialized in Materials Science, Catalysis, and Environmental Engineering.
Your task is to precisely extract material substitution logic, explicit dosages, and performance metrics from scientific text into a structured JSON format.

# ⚠️ MANDATORY PRE-FILTER: Relevance Check (Execute FIRST)

Before ANY extraction, you MUST determine the ROLE of the target compound "{target}" in this paper:

| Role | Description | Action |
|------|-------------|--------|
| **Process Improvement** | The target is STILL the raw material, but with improved catalyst/process (e.g., better catalyst for ethylbenzene → styrene) | ❌ STOP - Return `"status": "irrelevant"` |
| **Pollutant/VOC/Emission** | The target is an unwanted byproduct or environmental contaminant being measured or reduced | ❌ STOP - Return `"status": "irrelevant"` |
| **Structural/Molecular Sub** | The target is a sub-molecular structural moiety (e.g., "benzene ring") replaced in drug design / scaffold hopping. | ❌ STOP - Return `"status": "irrelevant"` |
| **Generic/Illustrative Mention** | The target is named only as one item in a GENERIC LIST of substances to avoid (e.g., "organic solvents such as benzene, toluene, chloroform…") and the study does NOT specifically design an experiment to replace `{target}` by name. The substitution target of the paper is a class of solvents, not `{target}` individually. | ❌ STOP - Return `"status": "irrelevant"` |
| **Solvent/Additive/Formulation** | The target is used AS a functional ingredient (solvent, coating, plasticizer) and the paper proposes a REPLACEMENT — AND `{target}` is either the PRIMARY named target or a major component being explicitly replaced (not merely listed in a generic category). | ✅ PROCEED with extraction |
| **Feedstock Substitution** | A DIFFERENT raw material replaces the target to produce the SAME end product (e.g., bio-ethanol replaces ethylbenzene for styrene production) | ✅ PROCEED with extraction |

**Decision Logic:**
1. If the paper improves how to CONVERT {target} (but {target} is STILL needed) → Irrelevant (process improvement)
2. If the paper measures {target} as a pollution/emission source → Irrelevant (it's a pollutant)
3. If the paper proposes a safer/greener chemical to REPLACE {target} in an application → Relevant (proceed)
4. If the paper proposes a DIFFERENT feedstock to produce the same product that {target} makes → Relevant (feedstock substitution)
5. If `{target}` is merely cited in a generic list of harmful solvents/chemicals as background justification, while the study's named substitution target is a broader CATEGORY (e.g., "volatile organic solvents", "BTX aromatics") and `{target}` is not the specific compound being replaced in the experimental protocol → Irrelevant (generic list mention)

**If Irrelevant, return IMMEDIATELY:**
```json
{{
  "status": "irrelevant",
  "reason": "[Target] is used as a [reactant|pollutant|measurement subject], not a material being substituted.",
  "detected_role": "reactant | pollutant | measurement_subject",
  "substitution_logic": null,
  "dosage_found": false,
  "explicit_dosages": null,
  "synthesis_conditions": null,
  "material_properties": null,
  "performance_metrics": null,
  "confidence": "high"
}}
```

**Only if Relevant, continue to extraction below.**

---

# Strict Extraction Constraints

## 1. Contextual Substitution Logic (`substitution_logic`)

**Domain Identification Rule:**
- FIRST, identify the paper's domain by reading the abstract and title.
- THEN, adapt the `target_problem` accordingly:

| Domain | target_problem Examples | Notes |
|--------|------------------------|-------|
| **Catalysis / Chemical Synthesis** | "high energy consumption", "thermodynamic limitations", "catalyst deactivation by coke", "low selectivity", "harsh reaction conditions" | Chemicals like ethylbenzene are REACTANTS, NOT pollutants |
| **Environmental Remediation** | "pollutant emissions", "VOC reduction", "wastewater treatment", "air quality improvement" | The target chemical IS the pollutant |
| **Materials Engineering** | "poor mechanical properties", "limited durability", "high production cost", "thermal instability" | Focus on material performance gaps |

**CRITICAL:** NEVER blindly default to "emission/pollution reduction" for all papers. Read the context!

## 2. Zero-Inference Dosage Extraction (`explicit_dosages`)

**Extraction Rules:**
- Extract ONLY explicitly stated material compositions, doping levels, or synthesis ratios
- **CRITICAL RESTRICTION ON `material`**: The extracted `material` MUST be either the `{alternative}` or the `{target}`. NEVER extract dosages for auxiliary chemicals, buffers, salts, solvents, or dye assistants. If the text only has numbers for these unrelated chemicals, treat it as missing dosage!
- **FORBIDDEN — Synthesis-of-alternative dosages**: If the paper describes HOW TO MAKE the alternative compound (e.g., synthesis reaction of the proposed replacement), the raw materials and molar ratios used in THAT synthesis are NOT the dosage of the alternative replacing `{target}`. Those must go into `synthesis_conditions`, NOT `explicit_dosages`. Dosages in `explicit_dosages` must reflect the USAGE or LOADING of the alternative in the APPLICATION context where it replaces `{target}`.
- Typical sources: Abstract, "Materials and Methods", "Experimental Section", Tables
- Each dosage MUST include complete physical units

**CRITICAL: Baseline Comparison Required**
- ALWAYS extract BOTH the alternative material's dosage AND the target compound's baseline dosage when available
- The baseline (target compound) dosage is essential for calculating relative performance improvement
- Use `"role": "baseline"` for the target compound and `"role": "alternative"` for the replacement
- **SCAN THE ENTIRE TEXT** for comparison statements like "X contained ~0.5 mmol/g whereas Y displayed ~1.9 mmol/g"
- When multiple materials are compared in the SAME PARAGRAPH, extract ALL their dosages
- If a numerical value exists for the target compound (even if lower/worse), you MUST capture it
- DO NOT mark baseline value as "not explicitly stated" if the text contains a number like "~0.5 mmol/g"

**ENUM STRICT RESTRICTION for `ratio_type`:**
Choose ONLY from this allowed list:
- `"wt%"` - weight percentage
- `"at%"` - atomic percentage  
- `"vol%"` - volume percentage
- `"mass_ratio"` - mass ratio (e.g., 1:1, 3:6:1)
- `"molar_ratio"` - molar ratio
- `"catalyst_loading"` - catalyst amount (mg, g, wt% on support)
- `"concentration"` - solution concentration (mM, M, mg/L, mol/L)

**FORBIDDEN ratio_types:** Do NOT use legacy terms like "binder_over_sand", "hardener_over_binder", "reaction_concentration" unless the paper explicitly uses these exact terms.

**ZERO INFERENCE POLICY:**
- NEVER calculate, multiply, or guess mass/volume based on assumptions
- If no exact number is found in the text, set `dosage_found` to false
- Do NOT derive values from percentages × totals (e.g., no "110 kg × 2% = 2.2 kg")

## 3. Dynamic Performance Metrics (`performance_metrics`)

**Mutually Exclusive Rule:**
NEVER put the following into `explicit_dosages`:
- Rates with time units (mmol/g/h, h⁻¹, mol·L⁻¹·s⁻¹)
- Yield percentages
- Conversion rates
- Selectivity values
- TOF (turnover frequency)
- kcat values
- Removal/reduction percentages

These MUST go into `performance_metrics` array.

**Dynamic Naming:**
- Define `metric_name` based on what the paper evaluates
- Examples: "styrene_selectivity", "ethylbenzene_conversion", "BTEX_reduction", "specific_activity", "TOF"

**Comparison Baseline:**
- If the paper provides a reference point (e.g., "compared to pristine catalyst 12.1%"), capture it
- If no baseline exists, set `comparison_baseline` to null

## 4. Synthesis Conditions (`synthesis_conditions`)

**Purpose:** Capture process parameters used during material preparation
**Allowed Data Types:**
- Temperature (°C, K)
- Pressure (atm, bar, Pa)
- Time/Duration (h, min, s)
- Atmosphere (N2, Ar, air, vacuum)
- pH values
- Calcination/Annealing conditions
- CVD/ALD cycle counts

**CRITICAL:** These are NOT dosages - they describe HOW the material was made, not its composition.

## 5. Material Properties (`material_properties`)

**Purpose:** Capture physical/structural characteristics of the synthesized material
**Allowed Data Types:**
- Layer count (e.g., "2-5 layers")
- Particle size (nm, μm)
- Surface area (m²/g, BET)
- Pore size/volume
- Crystallinity
- Morphology descriptors
- Thickness
- **For surfactants:** Critical micelle concentration (CMC), surface tension at CMC (γCMC), Krafft point — these are intrinsic physical properties of the surfactant, NOT application dosages. Place them in `material_properties`, NOT in `explicit_dosages`. An application dosage would be the actual use concentration in a formulation (e.g., wt% in a foam, mg/L in a cleaning bath).
- **Elemental / molecular composition:** wt% or mol% describing the elemental content of a synthesized material (e.g., "fluorine content 19.3 wt%", "fluorous content 51 wt%", "carbon content 65 wt%") is a structural characteristic, NOT an application dosage. Place it in `material_properties`.
- **Measurement conditions ≠ application dosages:** A concentration used only for characterisation (e.g., "1 wt% aqueous solution for surface tension measurement", "0.1 mg/mL stock for spectroscopy") is NOT an application dosage. Do NOT place it in `explicit_dosages`. Report `insufficient_data` unless a separate, real-world use concentration is also stated.

**CRITICAL:** These describe WHAT the material IS, not its formulation ratio. An `explicit_dosage` must answer "how much of the alternative was added to achieve the desired function in a real application."

---

TARGET POLLUTANT/REACTANT: {target}
ALTERNATIVE SOLUTION: {alternative}
PAPER TITLE: {title}
DOI: {doi}

TEXT CONTENT:
{text}

---

Return JSON with this exact schema:
{{
  "substitution_logic": {{
    "target_problem": "the SPECIFIC technical pain point from THIS paper (domain-aware, NOT generic)",
    "traditional_material_replaced": "the conventional material/process being replaced",
    "alternative_solution": "the safer/better replacement material/process",
    "relationship_type": "catalyst_substitution | material_substitution | process_substitution | solvent_substitution | fuel_substitution"
  }},
  "dosage_found": true/false,
  "explicit_dosages": [
    {{
      "material": "specific material name as stated in paper",
      "value": "exact value with unit (e.g., '0.2 at%', '30 mg', '2.5 wt%')",
      "ratio_type": "wt% | at% | vol% | mass_ratio | molar_ratio | catalyst_loading | concentration",
      "role": "baseline | alternative",
      "evidence_location": "Table X | Section Y.Z | Figure caption | exact quote",
      "context": "brief experimental context"
    }}
  ],
  "synthesis_conditions": [
    {{
      "parameter": "condition name (e.g., calcination_temperature, CVD_cycles, reaction_time, atmosphere)",
      "value": "exact value with unit (e.g., '800 °C', '12 cycles', '2 h', 'N2')",
      "evidence_location": "Section X | Table Y",
      "context": "brief description of the synthesis step"
    }}
  ],
  "material_properties": [
    {{
      "property": "property name (e.g., layer_count, particle_size, surface_area, pore_volume)",
      "value": "exact value with unit (e.g., '2-5 layers', '50 nm', '120 m²/g')",
      "evidence_location": "Section X | Figure Y | Table Z",
      "context": "measurement method or characterization technique if mentioned"
    }}
  ],
  "performance_metrics": [
    {{
      "metric_name": "dynamically determined (e.g., conversion_rate, selectivity, TOF, BTEX_reduction)",
      "metric_value": "exact value with unit as stated (e.g., '32.6%', '25.3 mmol·g⁻¹·h⁻¹', '>90%')",
      "comparison_baseline": "reference point if mentioned, or null"
    }}
  ],
  "confidence": "high | medium | low"
}}

**REMINDER:** Even if dosage_found is false, you MUST still extract all performance metrics from the abstract/conclusions.
"""

DOSAGE_INFERENCE_PROMPT = """You are a chemistry expert. The paper did NOT provide explicit dosage data. Your task is to determine if DETERMINISTIC CALCULATION is possible based on partial data in the text.

=== INFERENCE GUIDELINES (推算守則) ===

**1. ENTITY DECOUPLING PRINCIPLE (實體解耦原則)**
- NEVER equate the concentration/mass of the TARGET POLLUTANT (e.g., VOCs, ethylbenzene) with the dosage of the ALTERNATIVE SOLUTION (e.g., geopolymer, sodium silicate).
- The alternative material dosage can ONLY be calculated based on:
  * CARRIER weight (載體重量): e.g., sand mold weight, substrate mass
  * SOLVENT volume (溶劑體積): e.g., reaction medium volume
  * SUPPORT surface area (載體表面積): e.g., catalyst support area
- The pollutant emission level is an OUTPUT METRIC, not an input for dosage calculation.

**2. PARAMETER HARVESTING & PROVENANCE (參數收割與溯源)**
Before any calculation, you MUST extract these BASE VARIABLES from the text:
- Total carrier/substrate weight or volume (with evidence location)
- Percentage ratio stated in the paper (with evidence location)
- Any explicit mass or volume values (with evidence location)

Each extracted variable MUST include:
- Exact value with unit
- Source location: "Table X", "Section Y.Z", "line quote: '...'"

**3. DETERMINISTIC CALCULATION TRACE (確定性運算軌跡)**
- FORBIDDEN: Using unstated "industry rules of thumb" (e.g., "typically 10-25%")
- REQUIRED: All calculations must be shown as mathematical formulas
  Example: "110 kg (sand) × 1.6% (binder ratio from Table 3) = 1.76 kg"
- If a CRITICAL VARIABLE is missing (e.g., percentage given but no total weight), 
  you MUST abandon the calculation and set calculated_dosage to null.

=== STRICT CONSTRAINTS ===
1. DO NOT fabricate dosage values based on molecular weight ratios.
2. DO NOT use "typical concentrations" from external knowledge.
3. Only perform calculations when ALL required variables are explicitly stated.

TARGET POLLUTANT: {target}
ALTERNATIVE SOLUTION: {alternative}
PAPER TITLE: {title}

AVAILABLE TEXT:
{text}

REASONING FROM EARLIER ANALYSIS:
{reasoning}

=== WHAT TO REPORT ===
Look for these types of information:
- Base variables that could enable calculation
- Qualitative descriptors: "low concentration", "excess", "stoichiometric"
- Relative comparisons: "similar to", "higher than", "reduced by X%"
- References to other papers with specific dosage data

Return JSON:
{{
  "partial_data_found": true/false,
  "base_variables_extracted": [
    {{
      "variable_name": "e.g., total_sand_weight | binder_percentage | substrate_mass",
      "value": "exact value with unit",
      "evidence_location": "Table X | Section Y | exact quote"
    }}
  ],
  "calculation_attempted": true/false,
  "calculated_dosage": {{
    "formula": "mathematical expression showing the calculation",
    "result": "calculated value with unit (or null if calculation cannot be completed)",
    "variables_used": ["list of variable names used"],
    "missing_variables": ["list of variables needed but not found in text"]
  }},
  "qualitative_indicators": [
    {{
      "description": "exact quote or paraphrase from text",
      "implication": "what this suggests without numerical speculation"
    }}
  ],
  "referenced_studies": [
    {{
      "citation": "if paper references another study with dosage data",
      "relevance": "why this reference might be useful"
    }}
  ],
  "data_gaps": [
    "specific information that would be needed but is missing"
  ],
  "recommendation": "suggest_fulltext_review | suggest_cited_reference | insufficient_data"
}}
"""


# ============================================================
# Smart Chunking System
# ============================================================

CHUNK_RELEVANCE_PROMPT = """You are quickly scanning a text segment from a scientific paper.

Task: Rate if this segment contains dosage/concentration OR performance/reduction data about the alternative substance.

TARGET POLLUTANT: {target}
ALTERNATIVE SOLUTION: {alternative}

=== TEXT SEGMENT ===
{chunk}

=== SCORING CRITERIA ===
Score 0: No relevant information (theoretical discussion, literature review only)
Score 1: Mentions the alternative but no quantities
Score 2: Contains experimental methods with potential dosage hints
Score 3: Contains tables, formulas, specific numerical data, OR reduction/performance percentages
Score 4: Contains EXPLICIT dosage/concentration values OR emission reduction data (e.g., ">90% BTEX reduction")

Return JSON only:
{{"score": <0-4>, "reason": "<brief 10-word reason>"}}
"""


def split_text_into_chunks(text: str, chunk_size: int = 2500, overlap: int = 200) -> list[dict]:
    """
    Split text into overlapping chunks for relevance scoring.
    
    Args:
        text: Full text to split
        chunk_size: Target size of each chunk in characters
        overlap: Overlap between chunks to avoid cutting important content
    
    Returns:
        List of dicts with 'text', 'start', 'end' keys
    """
    chunks = []
    start = 0
    text_len = len(text)
    
    while start < text_len:
        end = min(start + chunk_size, text_len)
        
        # Try to end at a paragraph or sentence boundary
        if end < text_len:
            # Look for paragraph break first
            para_break = text.rfind('\n\n', start + chunk_size // 2, end)
            if para_break > start:
                end = para_break + 2
            else:
                # Look for sentence end
                sentence_end = text.rfind('. ', start + chunk_size // 2, end)
                if sentence_end > start:
                    end = sentence_end + 2
        
        chunks.append({
            'text': text[start:end],
            'start': start,
            'end': end
        })
        
        # Move to next chunk with overlap
        start = end - overlap if end < text_len else text_len
    
    return chunks


def score_chunk_relevance(
    chunk: str,
    target: str,
    alternative: str,
    client: "LLMClient"
) -> int:
    """
    Use LLM to quickly score a chunk's relevance for dosage information.
    Returns score 0-4.
    """
    prompt = CHUNK_RELEVANCE_PROMPT.format(
        target=target,
        alternative=alternative,
        chunk=chunk[:2000]  # Limit chunk size for scoring
    )
    
    try:
        content = client.chat_completion(
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=100,
            json_mode=True,
            model="gpt-4o-mini"  # Use fast model for scoring
        )
        
        if content:
            result = json.loads(content)
            return int(result.get("score", 0))
    except Exception:
        pass
    
    return 1  # Default score if scoring fails


def smart_chunk_selection(
    text: str,
    target: str,
    alternative: str,
    client: "LLMClient",
    max_output_len: int = 20000,
    min_score: int = 2
) -> str:
    """
    Intelligently select relevant chunks from long text.
    
    Strategy:
    1. If text is short enough, return as-is
    2. Split into chunks and score each
    3. Keep high-scoring chunks (score >= min_score)
    4. Always keep first chunk (abstract/intro) and last chunk (conclusions)
    5. Combine selected chunks up to max_output_len
    
    Args:
        text: Full text to process
        target: Target pollutant
        alternative: Alternative solution
        client: LLM client wrapper
        max_output_len: Maximum output text length
        min_score: Minimum relevance score to include a chunk
    
    Returns:
        Selected text with markers for truncated sections
    """
    # If text is short enough, return as-is
    if len(text) <= max_output_len:
        return text
    
    print(f"      [SMART CHUNK] Text length {len(text):,} > {max_output_len:,}, applying intelligent selection...")
    
    # Split into chunks
    chunks = split_text_into_chunks(text, chunk_size=2500, overlap=200)
    print(f"      [SMART CHUNK] Split into {len(chunks)} chunks, scoring relevance...")
    
    # Pattern-based pre-scorer: boost chunks that contain the alternative name
    # AND a numeric dosage indicator in close proximity.  This guards against
    # the LLM scorer underrating chunks that use chemical shorthand such as
    # "4 equivalents", "10 mol%", "2:1 ratio", etc.
    _DOSAGE_PATTERN = re.compile(
        r'\d+(\.\d+)?\s*'
        r'(equiv(alent)?s?|eq\.?|mmol|mol\s*%|wt\s*%|g/[lL]|mg/[lL]|'
        r'μ[mM]|n[mM]|m[mM]|[μμ]mol|[mM]ol/[lL]|'
        r'equivalents?\s+of|parts?)',
        re.IGNORECASE,
    )

    def _get_alt_variants(alt: str) -> list[str]:
        """Generate simple chemical-name variants for matching."""
        alt_lower = alt.lower().strip()
        variants = {alt_lower}

        suffixes = ["carbonate", "chloride", "alcohol", "oxide", "ester", "ether", "amine", "acid"]
        prefixes = ["diallyl", "allyl", "methyl", "ethyl", "propyl", "butyl", "vinyl", "phenyl"]

        # normalize hyphen/space variants
        variants.add(alt_lower.replace("-", " "))
        variants.add(alt_lower.replace(" ", ""))
        variants.add(alt_lower.replace("-", ""))

        # split known suffix only if it is at the end
        for suffix in suffixes:
            if alt_lower.endswith(suffix) and not alt_lower.endswith(f" {suffix}"):
                base = alt_lower[:-len(suffix)].strip("- ")
                if base:
                    variants.add(f"{base} {suffix}")

        # split known prefix only if directly attached
        for prefix in prefixes:
            if alt_lower.startswith(prefix) and not alt_lower.startswith(f"{prefix} "):
                rest = alt_lower[len(prefix):].strip("- ")
                if rest:
                    variants.add(f"{prefix} {rest}")

        # also generate hyphenated forms from spaced variants
        extra = set()
        for v in variants:
            if " " in v:
                extra.add(v.replace(" ", "-"))
        variants.update(extra)

        return sorted(v for v in variants if v)

    def _pattern_boost(chunk_text: str) -> int:
        """Return 3 if chunk has numeric dosage pattern near the alternative OR target (baseline)."""
        text_lower = chunk_text.lower()
        
        # Try all name variants for ALTERNATIVE
        for alt_variant in _get_alt_variants(alternative):
            for m in re.finditer(re.escape(alt_variant), text_lower):
                vicinity = text_lower[max(0, m.start() - 400): m.end() + 400]
                if _DOSAGE_PATTERN.search(vicinity):
                    return 3
        
        # ALSO check TARGET for baseline dosage comparison data
        for target_variant in _get_alt_variants(target):
            for m in re.finditer(re.escape(target_variant), text_lower):
                vicinity = text_lower[max(0, m.start() - 400): m.end() + 400]
                if _DOSAGE_PATTERN.search(vicinity):
                    return 3  # Baseline comparison data is equally important
        
        return 0

    # Score each chunk
    scored_chunks: list[tuple[int, int, dict]] = []  # (score, index, chunk)
    for i, chunk in enumerate(chunks):
        # Always include first and last chunks
        if i == 0 or i == len(chunks) - 1:
            score = 4  # Ensure first/last are always included
        else:
            pattern_score = _pattern_boost(chunk['text'])
            # If pattern_boost found dosage near alternative, prioritize this chunk highly
            # Score 5 ensures these chunks are selected before LLM-scored chunks
            if pattern_score >= 3:
                score = 5  # Highest priority - dosage data detected
            else:
                llm_score = score_chunk_relevance(chunk['text'], target, alternative, client)
                score = max(pattern_score, llm_score)

        scored_chunks.append((score, i, chunk))
    
    # Sort by score (descending) then by position (ascending) for tie-breaking
    scored_chunks.sort(key=lambda x: (-x[0], x[1]))
    
    # Select chunks until we reach max_output_len
    selected_indices: set[int] = set()
    current_len = 0
    
    for score, idx, chunk in scored_chunks:
        if score < min_score:
            continue
        
        chunk_len = len(chunk['text'])
        if current_len + chunk_len <= max_output_len:
            selected_indices.add(idx)
            current_len += chunk_len
    
    # Build output in original order
    output_parts: list[str] = []
    prev_end = 0
    
    for i, chunk in enumerate(chunks):
        if i in selected_indices:
            # Add truncation marker if there's a gap
            if chunk['start'] > prev_end + 100:  # Gap > 100 chars
                output_parts.append("\n\n[... section skipped (low relevance) ...]\n\n")
            output_parts.append(chunk['text'])
            prev_end = chunk['end']
    
    result = ''.join(output_parts)
    
    # Count how many chunks were kept
    high_score_count = sum(1 for s, _, _ in scored_chunks if s >= min_score)
    print(f"      [SMART CHUNK] Selected {len(selected_indices)}/{len(chunks)} chunks (score >= {min_score}), output: {len(result):,} chars")
    
    return result


# ============================================================
# LLM Functions
# ============================================================

def extract_dosage_from_text(
    text: str,
    title: str,
    doi: str,
    target: str,
    alternative: str,
    client: "LLMClient"
) -> dict[str, Any]:
    """Use LLM to extract explicit dosage information from text."""
    # Use smart chunking for long texts instead of simple truncation
    max_text_len = 20000
    if len(text) > max_text_len:
        text = smart_chunk_selection(
            text=text,
            target=target,
            alternative=alternative,
            client=client,
            max_output_len=max_text_len,
            min_score=2
        )
    
    prompt = DOSAGE_EXTRACTION_PROMPT.format(
        target=target,
        alternative=alternative,
        title=title,
        doi=doi,
        text=text
    )
    
    for attempt in range(MAX_RETRIES):
        try:
            content = client.chat_completion(
                messages=[{"role": "user", "content": prompt}],
                temperature=0.0,
                max_tokens=1000,
                json_mode=True
            )
            
            if content:
                return json.loads(content)
            return {"dosage_found": False, "dosages": []}
            
        except json.JSONDecodeError:
            continue
        except Exception as e:
            if attempt == MAX_RETRIES - 1:
                print(f"    [WARN] LLM extraction failed: {e}")
            
    return {"dosage_found": False, "dosages": []}


def infer_dosage(
    text: str,
    title: str,
    target: str,
    alternative: str,
    reasoning: str,
    client: "LLMClient"
) -> dict[str, Any]:
    """Use LLM to infer dosage when explicit values are not available."""
    # Use smart chunking for long texts
    max_text_len = 12000
    if len(text) > max_text_len:
        text = smart_chunk_selection(
            text=text,
            target=target,
            alternative=alternative,
            client=client,
            max_output_len=max_text_len,
            min_score=1  # Lower threshold for inference
        )
    
    prompt = DOSAGE_INFERENCE_PROMPT.format(
        target=target,
        alternative=alternative,
        title=title,
        text=text if text else "(No fulltext available - using abstract only)",
        reasoning=reasoning
    )
    
    for attempt in range(MAX_RETRIES):
        try:
            content = client.chat_completion(
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,  # Slightly higher for inference
                max_tokens=1200,
                json_mode=True
            )
            
            if content:
                return json.loads(content)
            return {}
            
        except json.JSONDecodeError:
            continue
        except Exception as e:
            if attempt == MAX_RETRIES - 1:
                print(f"    [WARN] LLM inference failed: {e}")
    
    return {}


# ============================================================
# Fulltext Quality Check
# ============================================================

_ERROR_PAGE_PATTERNS: list[str] = [
    "the requested url was rejected",
    "access denied",
    "403 forbidden",
    "404 not found",
    "page not found",
    "unauthorized",
    "your request has been blocked",
    "enable javascript",
    "please enable cookies",
    "robot or human",
    "captcha",
    "cloudflare",
    "just a moment",
    "checking your browser",
    "ddos protection",
    "too many requests",
]

_MIN_MEANINGFUL_CHARS = 300


def _is_meaningful_fulltext(text: str, title: str = "") -> bool:
    """
    Return True if ``text`` looks like genuine scientific paper content.
    Return False if it appears to be an error page, access-denial message,
    or otherwise too short to be useful.
    """
    stripped = text.strip()
    if len(stripped) < _MIN_MEANINGFUL_CHARS:
        return False
    lower = stripped.lower()
    for pattern in _ERROR_PAGE_PATTERNS:
        if pattern in lower:
            return False
    return True


# ============================================================
# Main Processing
# ============================================================

def process_paper_for_dosage(
    record: dict[str, Any],
    target: str,
    research_pdf_dir: Path,
    client: "LLMClient"
) -> dict[str, Any]:
    """Process a single paper to extract or infer dosage information."""
    result = record.copy()
    
    doi = record.get("doi", "")
    title = record.get("title", "")
    # Handle alternatives as list (from step03) or string (legacy step04)
    alternatives_raw = record.get("alternatives", "")
    if isinstance(alternatives_raw, list):
        alternative = ", ".join(alternatives_raw)
    else:
        alternative = str(alternatives_raw) if alternatives_raw else ""
    reasoning = record.get("reasoning", "")
    abstract = record.get("abstract", "")
    
    # Get download strategy from step03 result
    download_status = record.get("download_status", "")
    
    if not alternative:
        result["dosage_info"] = {
            "status": "no_alternative",
            "explicit_dosage": None,
            "inferred_dosage": None,
            "extraction_method": None
        }
        return result
        
    # Try to find an existing PDF mapping strictly in the folder regardless of SemanticScholar results
    # We still check if DOI can locate a local file.
    
    # Step 1: Try to get fulltext
    fulltext, source = get_fulltext(doi, research_pdf_dir)
    
    if source == "deleted_pdf":
        print(f"    [INFO] Fake PDF deleted for {doi}. Forcing a re-download attempt via step03...")
        text_to_analyze = ""
        try:
            from step03 import download_full_text
            dl_status = download_full_text(doi, title, research_pdf_dir)
            fulltext, source = get_fulltext(doi, research_pdf_dir)
            if source != "none" and fulltext.strip():
                text_to_analyze = fulltext
                download_status = dl_status or f"Downloaded {source} after deleting fake PDF"
            else:
                text_to_analyze = abstract
        except ImportError:
            text_to_analyze = abstract
            
    else:
        # Check if download was reported as failure in step03,
        # but we DO have a file in research_pdf. This means step 04 should 
        # definitely use the found file.
        if source != "none" and fulltext.strip():
            # It's an actual file and it has readable content. Let's make sure it's noted.
            if "Failed" in download_status or not download_status:
                download_status = f"Local check found valid {source}"
        
        # If we failed to get fulltext locally, we can also try to trigger a download
        # if the download_status says it failed, because maybe we can fetch it now.
        if source == "none" and ("Failed" in download_status or not download_status):
            print(f"    [INFO] No local file and step03 reported failure for {doi}. Retrying download now...")
            try:
                from step03 import download_full_text
                dl_status = download_full_text(doi, title, research_pdf_dir)
                fulltext, source = get_fulltext(doi, research_pdf_dir)
                if source != "none" and fulltext.strip():
                    download_status = dl_status or f"Downloaded {source} after missing local file"
            except ImportError:
                pass
                
        # Validate fulltext is meaningful scientific content (not an error page)
        if fulltext and not _is_meaningful_fulltext(fulltext, title):
            print(f"    [WARN] Fulltext for {doi} appears to be an error page or garbage; falling back to abstract.")
            fulltext = ""
            source = "none"

        text_to_analyze = fulltext if fulltext else abstract

        if not text_to_analyze:
            # If both fulltext and abstract are empty, try to trigger a download
            print(f"    [INFO] No text found for {doi}. Attempting to download via step03 functionality...")
            try:
                from step03 import download_full_text
                # Call step03's downloader mapping
                dl_status = download_full_text(doi, title, research_pdf_dir)
                
                # Check again after download attempt
                fulltext, source = get_fulltext(doi, research_pdf_dir)
                if source != "none" and fulltext.strip():
                    text_to_analyze = fulltext
                    download_status = dl_status or f"Downloaded {source} on demand"
            except ImportError:
                pass
            
    if not text_to_analyze:
        # Still nothing
        print(f"    [WARN] Completely empty text for {doi}, skipping extraction.")
        result["dosage_info"] = {
            "status": "extraction_error",
            "reason": "Text completely empty",
            "explicit_dosage": None,
            "inferred_dosage": None,
            "extraction_method": "none"
        }
        return result
        
    result["fulltext_source"] = source
    
    # Determine extraction method details
    if fulltext:
        # Parse download strategy from step03's download_status
        if "Elsevier" in download_status:
            extraction_method = "fulltext_pdf_elsevier"
        elif "Semantic Scholar" in download_status:
            extraction_method = "fulltext_pdf_semantic_scholar"
        elif "Springer" in download_status or "JATS" in download_status:
            extraction_method = "fulltext_xml_springer"
        elif "Playwright Scraper" in download_status:
            extraction_method = "fulltext_pdf_playwright_scraper"
        elif "Webpage Capture" in download_status:
            extraction_method = "fulltext_pdf_webpage_capture"
        elif "openAccessPdf" in download_status:
            extraction_method = "fulltext_pdf_open_access"
        elif source == "xml":
            extraction_method = "fulltext_xml"
        else:
            extraction_method = "fulltext_pdf"
    else:
        extraction_method = "abstract_only"
    
    # Step 2: Extract explicit dosage from text
    extraction_result = extract_dosage_from_text(
        text=text_to_analyze,
        title=title,
        doi=doi,
        target=target,
        alternative=alternative,
        client=client
    )
    
    # Handle "irrelevant" status from pre-filter
    if extraction_result.get("status") == "irrelevant":
        # If step03 confirmed alternatives exist in this paper, upgrade to insufficient_data
        # so the alternative names are preserved and Phase C can potentially trigger.
        step03_confirmed_alternatives = result.get("alternatives provided", "").lower() == "yes"
        detected_role = extraction_result.get("detected_role")
        if step03_confirmed_alternatives:
            known_alternatives = result.get("alternatives") or []
            dosage_info: dict[str, Any] = {
                "status": "insufficient_data",
                "reason": (
                    f"Paper confirmed alternatives exist but target is identified as "
                    f"{detected_role or 'pollutant/measurement subject'} — "
                    f"no application dosage found."
                ),
                "detected_role": detected_role,
                "known_alternatives": known_alternatives,
                "substitution_logic": None,
                "explicit_dosages": None,
                "synthesis_conditions": None,
                "material_properties": None,
                "performance_metrics": None,
                "partial_data": None,
                "text_source": source if fulltext else "abstract",
                "extraction_method": extraction_method,
                "download_strategy": download_status if download_status else "N/A",
                "confidence": "high"
            }
        else:
            dosage_info = {
                "status": "irrelevant",
                "reason": extraction_result.get("reason", "Target compound role mismatch"),
                "detected_role": detected_role,
                "substitution_logic": None,
                "explicit_dosages": None,
                "synthesis_conditions": None,
                "material_properties": None,
                "performance_metrics": None,
                "partial_data": None,
                "text_source": source if fulltext else "abstract",
                "extraction_method": extraction_method,
                "download_strategy": download_status if download_status else "N/A",
                "confidence": "high"
            }
        result["dosage_info"] = dosage_info
        return result
    
    # Classify extraction completeness (complete = has material dosage)
    completeness = classify_extraction_completeness(extraction_result)
    
    # Map completeness to status
    if completeness == "complete":
        status = "extracted"  # Full extraction with material dosage
    elif completeness == "partial_result_only":
        status = "partial_data"  # Has dosages but only result/output, not input
    else:
        status = "not_found"
    
    # New schema with substitution_logic support
    # Filter out synthesis-scale values and impurity concentrations before storing
    _raw_dosages = extraction_result.get("explicit_dosages") if extraction_result.get("dosage_found") else None
    _filtered_dosages = filter_functional_dosages(_raw_dosages) or None
    dosage_info: dict[str, Any] = {
        "status": status,
        "completeness": completeness,  # New field to indicate extraction quality
        "has_material_dosage": has_material_dosage(_filtered_dosages),
        "substitution_logic": extraction_result.get("substitution_logic"),
        "explicit_dosages": _filtered_dosages,
        "synthesis_conditions": extraction_result.get("synthesis_conditions"),
        "material_properties": extraction_result.get("material_properties"),
        "performance_metrics": extraction_result.get("performance_metrics"),
        "partial_data": None,
        "text_source": source if fulltext else "abstract",
        "extraction_method": extraction_method,
        "download_strategy": download_status if download_status else "N/A",
        "confidence": extraction_result.get("confidence", "low")
    }
    
    # Step 3: If no material dosage found, look for partial data (NO inference)
    if completeness != "complete":
        partial_result = infer_dosage(
            text=text_to_analyze,
            title=title,
            target=target,
            alternative=alternative,
            reasoning=reasoning,
            client=client
        )
        
        # Logic: prioritize completeness classification
        if completeness == "partial_result_only":
            # Has result dosages but missing material dosages → may need ESI
            dosage_info["status"] = "partial_data"
            dosage_info["missing"] = "material_dosage"
            dosage_info["recommendation"] = "suggest_esi_for_material_dosage"
            if partial_result:
                dosage_info["partial_data"] = partial_result
                # Upgrade to extracted if calculation is complete (result obtained, no missing variables)
                calc = partial_result.get("calculated_dosage", {})
                if (
                    partial_result.get("calculation_attempted")
                    and calc.get("result") is not None
                    and not calc.get("missing_variables")
                ):
                    dosage_info["status"] = "extracted"
        elif partial_result and partial_result.get("partial_data_found"):
            dosage_info["status"] = "partial_data"
            dosage_info["partial_data"] = partial_result
            # Upgrade to extracted if calculation is complete (result obtained, no missing variables)
            calc = partial_result.get("calculated_dosage", {})
            if (
                partial_result.get("calculation_attempted")
                and calc.get("result") is not None
                and not calc.get("missing_variables")
            ):
                dosage_info["status"] = "extracted"
        else:
            dosage_info["status"] = "insufficient_data"
            dosage_info["data_gaps"] = partial_result.get("data_gaps", []) if partial_result else []
            dosage_info["recommendation"] = partial_result.get("recommendation", "insufficient_data") if partial_result else "insufficient_data"
    
    result["dosage_info"] = dosage_info
    return result


def run_step04(
    input_file: Path,
    output_file: Path,
    target: str,
    research_pdf_dir: Optional[Path] = None,
    drop_empty: bool = False,
    cid: Optional[str] = None,
    final_dir: Optional[str] = None
) -> None:
    """Main execution function for Step 04."""
    print(f"\n{'='*60}")
    print(f"Step 04: Dosage Extraction & Inference")
    print(f"Target compound: {target}")
    print(f"Input file: {input_file}")
    print(f"{'='*60}\n")
    
    # Load step03 results
    if not input_file.exists():
        print(f"[ERROR] Input file not found: {input_file}")
        return
    
    with open(input_file, "r", encoding="utf-8") as f:
        step03_results = json.load(f)
    
    if not isinstance(step03_results, list):
        print("[ERROR] Invalid input format - expected list")
        return
    
    # Determine research_pdf directory
    if research_pdf_dir is None:
        research_pdf_dir = input_file.parent / "research_pdf"
    
    pdf_count = len(list(research_pdf_dir.glob('*'))) if research_pdf_dir.exists() else 0
    print(f"Research PDF directory: {research_pdf_dir}")
    print(f"PDF/XML files found: {pdf_count}")
    
    # Count records with alternatives
    # Handle both list (from step03) and string (legacy) formats
    def has_alternatives(r: dict) -> bool:
        alts = r.get("alternatives", "")
        if isinstance(alts, list):
            return len(alts) > 0
        return bool(alts)
    
    records_with_alts = [r for r in step03_results if has_alternatives(r)]
    print(f"Records with alternatives: {len(records_with_alts)}/{len(step03_results)}\n")
    
    if not records_with_alts:
        print("[INFO] No alternatives found - nothing to process")
        output_file.parent.mkdir(parents=True, exist_ok=True)
        with open(output_file, "w", encoding="utf-8") as f:
            json.dump({"no_alternatives": True}, f, ensure_ascii=False, indent=2)
        return
    
    # --- Dosage Extraction Caching System ---
    out_dir = output_file.parent
    cached_dosage_results = {}
    cache_hits = 0
    if out_dir.exists():
        import glob
        existing_files = list(out_dir.glob("step04_results*.json"))
        for cache_file in existing_files:
            try:
                with open(cache_file, "r", encoding="utf-8") as cache_f:
                    cached_data = json.load(cache_f)
                    if isinstance(cached_data, dict) and cached_data.get("no_alternatives"):
                        continue
                    if isinstance(cached_data, list):
                        for c_rec in cached_data:
                            # Verify this record has some dosage info
                            d_info = c_rec.get("dosage_info")
                            if d_info:
                                doi = str(c_rec.get("externalIds", {}).get("DOI", "")).strip()
                                title = str(c_rec.get("title", "") or c_rec.get("Article title", "")).strip()
                                
                                if doi: cached_dosage_results[doi] = c_rec
                                if title: cached_dosage_results[title] = c_rec
            except Exception as e:
                print(f"Warning: Failed to read cache {cache_file.name}: {e}")
                
    print(f"Loaded {len(cached_dosage_results)} cached Step 04 records from {out_dir.name}")
    # --- End Caching System ---

    # Initialize LLM client    # Initialize LLM client    # Initialize LLM client (supports both OpenAI and Gemini)
    client = LLMClient()
    
    # Process ONLY records with alternatives (skip the rest)
    results: list[dict[str, Any]] = []
    stats = {
        "extracted": 0, 
        "partial_data": 0, 
        "partial_with_calculation": 0,  # New: successful deterministic calculation
        "insufficient_data": 0, 
        "not_found": 0,
        "irrelevant": 0  # Papers where target is reactant/pollutant, not substituted material
    }
    
    update_interval = max(1, len(records_with_alts) // 100)  # Update every 1%
    with tqdm(total=len(records_with_alts), desc="Processing papers with alternatives", miniters=update_interval) as pbar:
        for record in records_with_alts:
            alt_raw = record.get("alternatives", "")
            if isinstance(alt_raw, list):
                alt = ", ".join(alt_raw)
            else:
                alt = str(alt_raw) if alt_raw else ""
            pbar.set_postfix_str(f"{alt[:25]}...")
            
            # Use Cache if available
            doi = str(record.get("externalIds", {}).get("DOI", "")).strip()
            title = str(record.get("title", "") or record.get("Article title", "")).strip()
            
            processed = None
            if doi and doi in cached_dosage_results:
                processed = cached_dosage_results[doi]
            elif title and title in cached_dosage_results:
                processed = cached_dosage_results[title]
                
            if processed and "dosage_info" in processed:
                # Merge cached dosage info into current record to keep other info up-to-date
                record["dosage_info"] = processed["dosage_info"]
                processed = record
                cache_hits += 1
            else:
                processed = process_paper_for_dosage(
                    record=record,
                    target=target,
                    research_pdf_dir=research_pdf_dir,
                    client=client
                )
            
            results.append(processed)
            
            # Update stats
            status = processed.get("dosage_info", {}).get("status", "not_found")
            if status in stats:
                stats[status] += 1
            
            # Track successful deterministic calculations
            if status == "partial_data":
                partial_data = processed.get("dosage_info", {}).get("partial_data", {})
                calc_result = partial_data.get("calculated_dosage", {})
                if calc_result and calc_result.get("result") is not None:
                    stats["partial_with_calculation"] += 1
            
            pbar.update(1)
    
    # Save results (only papers with alternatives)
    output_file.parent.mkdir(parents=True, exist_ok=True)
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    
    # Print summary
    print(f"\n{'='*60}")
    print(f"Step 04 Complete")
    print(f"{'='*60}")
    print(f"Papers with alternatives processed: {len(results)}")
    print(f"  - Explicit dosage extracted: {stats['extracted']}")
    print(f"  - Partial data found: {stats['partial_data']}")
    print(f"    └─ With deterministic calculation: {stats['partial_with_calculation']}")
    print(f"  - Insufficient data: {stats['insufficient_data']}")
    print(f"  - Irrelevant (reactant/pollutant): {stats['irrelevant']}")
    print(f"  - Not found: {stats['not_found']}")
    print(f"\nOutput saved to: {output_file}")
    
    # Save detailed summary (new schema with substitution_logic)
    # Generate summary filename based on output filename (e.g., step04_results_gemini.json -> step04_summary_gemini.json)
    output_stem = output_file.stem  # e.g., "step04_results_gemini"
    if output_stem.startswith("step04_results"):
        suffix = output_stem.replace("step04_results", "")  # e.g., "_gemini" or ""
        summary_filename = f"step04_summary{suffix}.json"
    else:
        summary_filename = "step04_summary.json"
    summary_file = output_file.parent / summary_filename
    output_stem = output_file.stem
    if output_stem.startswith("step04_results"):
        suffix = output_stem.replace("step04_results", "")
        usage_filename = f"step04_token_usage{suffix}.json"
    else:
        usage_filename = "step04_token_usage.json"
    usage_file = output_file.parent / usage_filename
    with open(usage_file, "w", encoding="utf-8") as f:
        json.dump(client.usage, f, indent=2, ensure_ascii=False)
    print(f"Token usage saved to: {usage_file}")

    summary = {
        "target": target,
        "total_records": len(results),
        "statistics": stats,
        "records_with_dosage": [
            {
                "title": r.get("title", "")[:80],
                "doi": r.get("doi"),
                "alternative": r.get("alternatives"),
                "reasoning": r.get("reasoning"),
                "dosage_status": r.get("dosage_info", {}).get("status"),
                "fulltext_source": r.get("fulltext_source"),
                "extraction_method": r.get("dosage_info", {}).get("extraction_method"),
                "download_strategy": r.get("dosage_info", {}).get("download_strategy"),
                "substitution_logic": r.get("dosage_info", {}).get("substitution_logic"),
                "explicit_dosages": r.get("dosage_info", {}).get("explicit_dosages"),
                "synthesis_conditions": r.get("dosage_info", {}).get("synthesis_conditions"),
                "material_properties": r.get("dosage_info", {}).get("material_properties"),
                "performance_metrics": r.get("dosage_info", {}).get("performance_metrics"),
                "partial_data": r.get("dosage_info", {}).get("partial_data"),
                "confidence": r.get("dosage_info", {}).get("confidence")
            }
            for r in results
            if r.get("dosage_info", {}).get("status") in ("extracted", "partial_data")
        ]
    }
    
    with open(summary_file, "w", encoding="utf-8") as f:
        json.dump(summary, f, ensure_ascii=False, indent=2)
    print(f"Summary saved to: {summary_file}")
    
    # If CID and final_dir provided, create final output file
    if cid and final_dir:
        from datetime import datetime
        final_output_dir = Path(final_dir)
        final_output_dir.mkdir(parents=True, exist_ok=True)
        date_str = datetime.now().strftime("%Y%m%d")
        final_filename = f"{date_str}_{cid}.json"
        final_file = final_output_dir / final_filename
        
        # Filter results if drop_empty is True
        final_results = results
        if drop_empty:
            final_results = [r for r in results if r.get("dosage_info", {}).get("status") in ("extracted", "partial_data")]
        
        with open(final_file, "w", encoding="utf-8") as f:
            json.dump(final_results, f, ensure_ascii=False, indent=2)
        print(f"Final output saved to: {final_file} ({len(final_results)} records)")


def parse_args() -> argparse.Namespace:
    """Parse command line arguments."""
    parser = argparse.ArgumentParser(
        description="Step 04: Dosage Extraction & Inference from same paper"
    )
    parser.add_argument(
        "--input_file",
        required=True,
        help="Input JSON file (step03 results)"
    )
    parser.add_argument(
        "--output_file",
        required=True,
        help="Output JSON file path"
    )
    parser.add_argument(
        "--target",
        required=True,
        help="Target compound name (the original toxic chemical)"
    )
    parser.add_argument(
        "--research_pdf_dir",
        help="Directory containing downloaded PDFs/XMLs (default: same as input_file parent/research_pdf)"
    )
    parser.add_argument(
        "--drop_empty",
        action="store_true",
        help="Drop records with no alternatives found"
    )
    parser.add_argument(
        "--cid",
        help="PubChem Compound ID for the target compound"
    )
    parser.add_argument(
        "--final_dir",
        help="Directory for final output files (named by date_cid)"
    )
    return parser.parse_args()


def main() -> None:
    """Main entry point."""
    args = parse_args()
    
    research_pdf_dir = Path(args.research_pdf_dir) if args.research_pdf_dir else None
    
    run_step04(
        input_file=Path(args.input_file),
        output_file=Path(args.output_file),
        target=args.target,
        research_pdf_dir=research_pdf_dir,
        drop_empty=args.drop_empty,
        cid=args.cid,
        final_dir=args.final_dir
    )


if __name__ == "__main__":
    main()
